#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Jun 14 12:04:38 2019

@author: DannySwift
"""

import pandas as pd
import yaml
from datetime import datetime
from docx import Document
from simple_salesforce import Salesforce

#%%
config = yaml.safe_load(open('config.yml'))
sf = Salesforce(username=config['username'], 
                password=config['password'], 
                security_token=config['security_token'])

df = pd.read_csv(r'Downloads/GalaThanks.csv')
date = datetime.today().strftime('%B %d')

#%%
def format_doc(name, amount, date, don_date, tix, email):
    doc = Document(r'Downloads/Thanks.docx')
    for p in doc.paragraphs:
        if 'TODAYS DATE' in p.text:
            p.text = p.text.replace('TODAYS DATE', date)
        if 'NAME' in p.text:
            p.text = p.text.replace('NAME', name)
        if 'ADDRESS' in p.text:
            q = sf.query("SELECT MailingAddress FROM Contact WHERE Name = '{}'"
                         .format(name))
            try:
                address = ''.join([q['records'][0]['MailingAddress']['street'],
                                   '\n',
                                   q['records'][0]['MailingAddress']['city'],
                                   ', ',
                                   q['records'][0]['MailingAddress']['state'],
                                   ', ',
                                   q['records'][0]['MailingAddress']['postalCode']])
                p.text = p.text.replace('ADDRESS', address)
            except IndexError:
                p.text = p.text.replace('ADDRESS', email)
        if '---------' in p.text:
            p.text = p.text.replace('---------', amount)
        if 'XXX' in p.text:
            p.text = p.text.replace('XXX', name)
        if '[date]' in p.text:
            p.text = p.text.replace('[date]', don_date)
        if 'TIXVAL' in p.text:
            inline = p.runs
            for i in inline:
                if 'TIXVAL' in i.text:
                    i.text = i.text.replace('TIXVAL', str(int(amount) - int(tix)))
        if 'DEDVAL' in p.text:
            inline = p.runs
            for i in inline:
                if 'DEDVAL' in i.text:
                    i.text = i.text.replace('DEDVAL', tix)
    
    return(doc)

#%%
for _, row in df.iterrows():
    if row['AMOUNT TICKET'] != '-':
        if row['DATE'] != 'has not paid':
            new_doc = format_doc(row['NAME'], row['AMOUNT DONATION'], date, 
                                 row['DATE'], row['AMOUNT TICKET'], row['EMAIL'])
            new_doc.save(r'Documents/{}.docx'.format(row['NAME']))
            
#%%
def format_doc_pink(name, amount, date, don_date, email):
    doc = Document(r'Downloads/ThanksNoTix.docx')
    for p in doc.paragraphs:
        if 'TODAYS DATE' in p.text:
            p.text = p.text.replace('TODAYS DATE', date)
        if 'NAME' in p.text:
            p.text = p.text.replace('NAME', name)
        if 'ADDRESS' in p.text:
            q = sf.query("SELECT MailingAddress FROM Contact WHERE Name = '{}'"
                         .format(name))
            try:
                address = ''.join([q['records'][0]['MailingAddress']['street'],
                                   '\n',
                                   q['records'][0]['MailingAddress']['city'],
                                   ', ',
                                   q['records'][0]['MailingAddress']['state'],
                                   ', ',
                                   q['records'][0]['MailingAddress']['postalCode']])
                p.text = p.text.replace('ADDRESS', address)
            except IndexError:
                p.text = p.text.replace('ADDRESS', email)
        if '---------' in p.text:
            p.text = p.text.replace('---------', amount)
        if 'XXX' in p.text:
            p.text = p.text.replace('XXX', name.split(' ')[0])
        if '[date]' in p.text:
            p.text = p.text.replace('[date]', don_date)
        if 'DEDVAL' in p.text:
            inline = p.runs
            for i in inline:
                if 'DEDVAL' in i.text:
                    i.text = i.text.replace('DEDVAL', amount)
    
    return(doc)

#%%
for _, row in df.dropna(axis=0, subset=['DATE']).iterrows():
    if row['AMOUNT TICKET'] == '-':
        if row['AMOUNT DONATION'] != '-':
            new_doc = format_doc_pink(row['NAME'], row['AMOUNT DONATION'], 
                                      date, row['DATE'], row['EMAIL'])
            new_doc.save(r'Documents/{}.docx'.format(row['NAME']))
